#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""สร้างเอกสาร .docx อธิบายเงื่อนไขการกรองและการคำนวณของ Dashboard สภาวะช่องปาก"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, "/sessions/pensive-ecstatic-wozniak/mnt/.claude/skills/thai-official-document/scripts")
try:
    from thai_wrap import get_tokenizer, insert_zwsp
    TOK = get_tokenizer()[0]                               # (fn, engine_name)
except Exception:                                          # pragma: no cover
    TOK = None
    def insert_zwsp(t, tokenizer=None):
        return t

# หน้าจอของ Windows ใช้รหัสอักขระเดิม ทำให้พิมพ์ภาษาไทยแล้วโปรแกรมพัง จึงบังคับเป็น UTF-8
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

BASE = Path(__file__).resolve().parent.parent
OUT_DIR = BASE / "ผลลัพธ์"
DOCS = BASE / "docs"
FONT = "TH Sarabun New"
TEAL = RGBColor(0x0F, 0x76, 0x6E)
GREY = RGBColor(0x47, 0x55, 0x69)


def zw(text):
    return insert_zwsp(text, tokenizer=TOK) if TOK else text


def style_run(run, size=16, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONT)
    szcs = rpr.makeelement(qn("w:szCs"), {qn("w:val"): str(int(size * 2))})
    rpr.append(szcs)
    if bold:
        rpr.append(rpr.makeelement(qn("w:bCs"), {}))


def para(doc, text="", size=16, bold=False, color=None, align=None, space_after=4,
         space_before=0, indent=None, wrap=True, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align is not None:
        p.alignment = align
    if text:
        style_run(p.add_run(zw(text) if wrap else text), size, bold, color, italic)
    return p


def h1(doc, text):
    para(doc, text, size=19, bold=True, color=TEAL, space_before=14, space_after=6, wrap=False)


def h2(doc, text):
    para(doc, text, size=17, bold=True, space_before=10, space_after=4, wrap=False)


def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.7)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run("•  " if level == 0 else "–  "), 16)
    style_run(p.add_run(zw(text)), 16)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(12)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.makeelement(qn("w:rFonts"), {})
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), "Consolas")
    rpr.insert(0, rf)
    r.font.color.rgb = GREY
    return p


def table(doc, rows, widths, header=True, sizes=14, aligns=None):
    t = doc.add_table(rows=0, cols=len(widths))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].width = Cm(widths[j])
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            is_head = header and i == 0
            if aligns and not is_head:
                p.alignment = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER,
                               "r": WD_ALIGN_PARAGRAPH.RIGHT}[aligns[j]]
            elif is_head:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            txt = str(val)
            mono = txt.startswith("`") and txt.endswith("`")
            if mono:
                txt = txt[1:-1]
                r = p.add_run(txt)
                r.font.name = "Consolas"
                r.font.size = Pt(11.5)
                rpr = r._element.get_or_add_rPr()
                rf = rpr.makeelement(qn("w:rFonts"), {})
                for a in ("w:ascii", "w:hAnsi", "w:cs"):
                    rf.set(qn(a), "Consolas")
                rpr.insert(0, rf)
            else:
                style_run(p.add_run(zw(txt)), sizes, bold=is_head)
            if is_head:
                shd = cells[j]._tc.get_or_add_tcPr().makeelement(
                    qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "E8F3F1"})
                cells[j]._tc.get_or_add_tcPr().append(shd)
    return t


# ---------------------------------------------------------------- build
def build():
    s = json.loads((OUT_DIR / "summary.json").read_text(encoding="utf-8"))
    agg = {a: {} for a in s["ageOrder"]}
    for r in s["records"]:
        for k, v in r.items():
            if k in ("age", "hoscode"):
                continue
            agg[r["age"]][k] = agg[r["age"]].get(k, 0) + v
    L = s["ageLabel"]
    f = lambda n: f"{n:,.0f}"
    pc = lambda a, b: "-" if not b else f"{a / b * 100:.1f}"

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(2.0)
    sec.left_margin, sec.right_margin = Cm(2.5), Cm(2.0)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(16)
    st.element.rPr.rFonts.set(qn("w:cs"), FONT)

    # ---------- ปก ----------
    para(doc, "เอกสารอธิบายวิธีการคัดกรองข้อมูลและเงื่อนไขการคำนวณ",
         size=24, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, wrap=False)
    para(doc, "Dashboard วิเคราะห์ผลการสำรวจสภาวะช่องปาก กลุ่มอายุ 3, 6, 12 ปี และ 60 ปีขึ้นไป",
         size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, f"ข้อมูลจาก HDC Data Exchange · ปีงบประมาณ {s['fiscalYear']} · "
              f"จัดทำเมื่อ {s['generatedAt']}",
         size=15, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    para(doc, "เอกสารฉบับนี้อธิบายว่า Dashboard คัดกรองข้อมูลจากไฟล์ต้นทางอย่างไร และแต่ละตารางในรายงาน "
              "ใช้เงื่อนไขใดในการนับ เพื่อให้ผู้สนใจตรวจสอบ ทำซ้ำ หรือนำไปปรับใช้กับพื้นที่ของตนเองได้ "
              "เงื่อนไขทั้งหมดอ้างอิงเอกสาร “ตารางรายงานและเงื่อนไขการคำนวณ.docx” และคำนิยามคอลัมน์จากไฟล์ dentalfile.csv",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

    # ---------- 1 ----------
    h1(doc, "๑. แหล่งข้อมูลและการรวมไฟล์")
    para(doc, "ข้อมูลต้นทางเป็นไฟล์ Excel ที่ส่งออกจากระบบ HDC (Data Exchange การตรวจฟัน) "
              "จัดเก็บแยกโฟลเดอร์ตามกลุ่มอายุ ดังนี้", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    code(doc, "data/3ปี/  data/6ปี/  data/12ปี/  data/60ปี/")

    para(doc, "รายงานที่ต้องส่งออกจากระบบ HDC", bold=True, space_before=6, space_after=3)
    table(doc, [
        ["โฟลเดอร์", "ชื่อรายงานใน HDC"],
        ["`data/3ปี/`", "18.2 OHSP ร้อยละเด็กกลุ่มอายุ 3 ปีมีฟันผุในฟันน้ำนม"],
        ["`data/6ปี/`", "18.3 OHSP ร้อยละเด็กกลุ่มอายุ 6 ปีมีฟันผุในฟันแท้"],
        ["`data/12ปี/`", "18.5 OHSP ร้อยละเด็กกลุ่มอายุ 12 ปีมีฟันผุในฟันแท้"],
        ["`data/60ปี/`", "18.8 OHSP ร้อยละของกลุ่มก่อนวัยสูงอายุ ที่มีฟันแท้ใช้งานได้ไม่น้อยกว่า 20 ซี่"],
    ], [3.4, 13.1], sizes=13, aligns=["l", "l"])
    para(doc, "ทุกรายการต้องส่งออกด้วยปุ่ม Data Exchange ซึ่งให้ข้อมูลรายบุคคลพร้อมคอลัมน์ผลตรวจฟัน "
              "ไม่ใช่ปุ่มส่งออกตารางสรุป",
         size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=3, space_after=6)
    bullet(doc, "แต่ละโฟลเดอร์มีได้มากกว่าหนึ่งไฟล์ กรณีข้อมูลระดับจังหวัดที่ระบบแบ่งส่งออกทีละไม่เกิน 100,000 รายการ "
                "(เช่น กลุ่ม 60 ปีขึ้นไป มี 3 ไฟล์) ทุกไฟล์ในโฟลเดอร์เดียวกันจะถูกนำมาต่อกันแบบ append")
    bullet(doc, "ทุกไฟล์ใช้ชีตชื่อ Data และมีโครงสร้าง 43 คอลัมน์เหมือนกัน สคริปต์จะตรวจสอบและหยุดทำงานหากพบโครงสร้างไม่ตรงกัน")
    bullet(doc, "ระหว่างรวมไฟล์ จะเพิ่มคอลัมน์ agesurvey บันทึกชื่อโฟลเดอร์ต้นทาง เพื่อระบุว่าแต่ละรายการเป็นของกลุ่มอายุใด "
                "(ไม่ใช้คอลัมน์ denttype ในการแบ่งกลุ่มอายุ เพราะพบค่าที่ไม่ตรงกับกลุ่มปะปนอยู่)")
    bullet(doc, "ผลลัพธ์บันทึกเป็นไฟล์เดียวชื่อ combined.csv")
    bullet(doc, "ค่าว่างในไฟล์ต้นทางเก็บเป็นข้อความ \"<NA>\" ระบบจะแปลงเป็นค่าว่างก่อนคำนวณทุกครั้ง")
    bullet(doc, "รหัสหน่วยบริการ (hoscode) เชื่อมกับไฟล์ hospitals.csv เพื่อดึงชื่อหน่วยบริการ อำเภอ และจังหวัด "
                "สำหรับใช้เป็นตัวกรอง (ไฟล์นี้เข้ารหัสอักขระแบบ cp874 / TIS-620)")

    para(doc, "จำนวนรายการที่รวมได้", bold=True, space_before=8, space_after=3)
    rows = [["กลุ่มอายุ", "จำนวนไฟล์", "จำนวนรายการทั้งหมด (record)"]]
    nfile = {"3ปี": 1, "6ปี": 1, "12ปี": 1, "60ปี": 3}
    for a in s["ageOrder"]:
        rows.append([L[a] + " ปี", nfile.get(a, 1), f(agg[a]["total"])])
    rows.append(["รวม", sum(nfile.values()), f(sum(agg[a]["total"] for a in s["ageOrder"]))])
    table(doc, rows, [5.0, 4.0, 6.5], aligns=["l", "c", "r"])

    # ---------- 2 ----------
    h1(doc, "๒. การคัดกรองข้อมูลก่อนวิเคราะห์")
    para(doc, "ข้อมูลต้นทางเป็นรายชื่อประชากรทั้งหมดในกลุ่มอายุ ไม่ใช่เฉพาะผู้ที่มารับบริการ "
              "การวิเคราะห์จึงคัดกรองเป็นสามชั้น ดังนี้", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    h2(doc, "ชั้นที่ ๑ – ประชากรทั้งหมด (total)")
    para(doc, "นับทุกรายการในไฟล์ ใช้เป็นตัวหารของ “ร้อยละความครอบคลุมการตรวจ” ที่แสดงบนการ์ดสรุปด้านบนของ Dashboard",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    h2(doc, "ชั้นที่ ๒ – ผู้ที่ได้รับการตรวจฟันในปีงบประมาณที่วิเคราะห์ (examined)")
    bullet(doc, "ต้องมีวันที่ให้บริการ คือ date_serv ไม่เป็นค่าว่าง")
    bullet(doc, "และวันที่ตรวจต้องอยู่ในปีงบประมาณเดียวกัน (1 ตุลาคม ถึง 30 กันยายน)")
    para(doc, "ปีงบประมาณที่ใช้วิเคราะห์กำหนดอัตโนมัติจากวันที่ตรวจล่าสุดในข้อมูล ตามกฎ "
              "หากเดือนของวันที่ล่าสุดเป็นตุลาคมถึงธันวาคม ให้เป็นปีงบประมาณถัดไป มิฉะนั้นเป็นปีเดียวกัน",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=3)
    table(doc, [
        ["ตัวอย่างวันที่ตรวจล่าสุด", "ปีงบประมาณที่ใช้", "ช่วงวันที่ที่นับ"],
        ["7 กรกฎาคม 2569", "2569", "1 ต.ค. 2568 – 30 ก.ย. 2569"],
        ["7 ตุลาคม 2569", "2570", "1 ต.ค. 2569 – 30 ก.ย. 2570"],
    ], [6.0, 4.0, 6.5], aligns=["l", "c", "l"])
    para(doc, f"ข้อมูลชุดปัจจุบันมีวันที่ตรวจล่าสุด {s['latestServe']} จึงวิเคราะห์ปีงบประมาณ {s['fiscalYear']} "
              f"(ค.ศ. {s['fyStart']} ถึง {s['fyEnd']}) "
              f"รายการที่ตรวจนอกช่วงนี้ {f(sum(agg[a]['examinedOther'] for a in s['ageOrder']))} รายการ "
              "ยังนับอยู่ในประชากรทั้งหมด แต่ไม่นับเป็นผู้ได้รับการตรวจ",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=4)

    h2(doc, "ชั้นที่ ๓ – ผู้ที่ข้อมูลผ่านเกณฑ์คุณภาพ (quality)")
    para(doc, "ตารางที่ 2 ถึง 7 ทุกตารางใช้กลุ่มนี้เป็นฐานการคำนวณ เงื่อนไขมีสองส่วนที่ต้องเป็นจริงพร้อมกัน",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    bullet(doc, "ผู้ตรวจต้องเป็นทันตแพทย์หรือทันตาภิบาล คือ providertype เป็น 02 หรือ 06")
    bullet(doc, "จำนวนฟันต้องอยู่ในพิสัยที่สมเหตุสมผลตามกลุ่มอายุ ดังตารางต่อไปนี้")
    table(doc, [
        ["กลุ่มอายุ", "เงื่อนไขพิสัยจำนวนฟัน (ต้องเป็นจริงทุกข้อ)"],
        ["3 ปี", "`dteeth` อยู่ระหว่าง 1 ถึง 20  และ  `dcaries + dfilling + dextract ≤ 20`  และ  `dteeth + dextract = 20`"],
        ["6 ปี", "`pteeth` อยู่ระหว่าง 1 ถึง 12  และ  `pfilling + pextract + pcaries ≤ 12`  และ  `pcaries + pfilling ≤ pteeth`"],
        ["12 ปี", "`pteeth` อยู่ระหว่าง 1 ถึง 28  และ  `pfilling + pextract + pcaries ≤ 28`  และ  `pcaries + pfilling ≤ pteeth`"],
        ["60+ ปี", "`pteeth` อยู่ระหว่าง 0 ถึง 32  และ  `permanent_permanent + permanent_prosthesis + prosthesis_prosthesis` "
                   "อยู่ระหว่าง 0 ถึง 10  และ  `pteeth + pextract > 0`"],
    ], [3.0, 13.5], sizes=13, aligns=["c", "l"])
    para(doc, "หมายเหตุ หากคอลัมน์ที่ใช้ตรวจสอบเป็นค่าว่าง ถือว่าไม่ผ่านเกณฑ์ และคอลัมน์ result ที่ติดมากับไฟล์ HDC "
              "ไม่ถูกนำมาใช้ เพราะระบบคำนวณเกณฑ์คุณภาพเองตามเอกสารข้างต้น",
         size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=3)

    para(doc, "ผลการคัดกรองของข้อมูลชุดปัจจุบัน", bold=True, space_before=8, space_after=3)
    rows = [["กลุ่มอายุ", "ประชากรทั้งหมด", "ตรวจนอกปีงบ", "ตรวจในปีงบ", "ร้อยละที่ตรวจ",
             "ผ่านเกณฑ์คุณภาพ", "ร้อยละผ่านเกณฑ์"]]
    for a in s["ageOrder"]:
        t = agg[a]
        rows.append([L[a] + " ปี", f(t["total"]), f(t["examinedOther"]), f(t["examined"]),
                     pc(t["examined"], t["total"]), f(t["quality"]), pc(t["quality"], t["examined"])])
    tt = {k: sum(agg[a][k] for a in s["ageOrder"]) for k in ("total", "examined", "examinedOther", "quality")}
    rows.append(["รวม", f(tt["total"]), f(tt["examinedOther"]), f(tt["examined"]),
                 pc(tt["examined"], tt["total"]), f(tt["quality"]), pc(tt["quality"], tt["examined"])])
    table(doc, rows, [2.6, 2.7, 2.3, 2.3, 2.2, 2.4, 2.3], sizes=13,
          aligns=["l", "r", "r", "r", "r", "r", "r"])

    # ---------- 3 ----------
    h1(doc, "๓. คำนิยามที่ใช้ร่วมกันในหลายตาราง")
    h2(doc, "๓.๑ ชุดฟันที่ใช้นับตามกลุ่มอายุ")
    bullet(doc, "อายุ 3 ปี ใช้ฟันน้ำนม คือคอลัมน์ dteeth, dcaries, dfilling, dextract")
    bullet(doc, "อายุ 6, 12 และ 60 ปีขึ้นไป ใช้ฟันแท้ คือคอลัมน์ pteeth, pcaries, pfilling, pextract")

    h2(doc, "๓.๒ การแปลงค่าสภาวะปริทันต์ (GumStatus)")
    para(doc, "คอลัมน์ gum เก็บผลตรวจ 6 ส่วนของช่องปาก (sextant) เป็นข้อความ 6 หลัก แต่ละหลักมีค่า "
              "0 = ปกติ, 1 = เหงือกอักเสบ, 2 = มีหินน้ำลาย, 3 = ปริทันต์อักเสบหรือมีฟันโยก, 9 = ไม่มีฟันหรือตรวจไม่ได้",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    para(doc, "ขั้นที่ 1 ทำความสะอาดค่า", bold=True, space_before=4, space_after=2)
    bullet(doc, "ค่าที่สั้นกว่า 6 หลัก ตีความว่าเป็นค่าเดียวกันทั้งปาก เช่น \"9\" หมายถึง 999999 และ \"0\" หมายถึง 000000")
    bullet(doc, "เลข 4 และ 5 ซึ่งมาจากรหัส CPI ชุดเดิม แปลงเป็น 3 (ปริทันต์อักเสบ)")
    bullet(doc, "ค่าว่าง ค่าที่ไม่ใช่ตัวเลข หรือความยาวผิดรูปแบบ ไม่นับในตารางที่ 4 และ 5 และรายงานแยกไว้ใต้ตาราง")
    para(doc, "ขั้นที่ 2 สรุปเป็นสภาวะรายบุคคล", bold=True, space_before=4, space_after=2)
    bullet(doc, "ถ้าทุก sextant มีค่า 9 ให้ GumStatus = 9 (ตรวจไม่ได้)")
    bullet(doc, "กรณีอื่น GumStatus = ค่าสูงสุดของ sextant ที่มีค่า 0 ถึง 3 เท่านั้น เช่น 900029 ได้ GumStatus = 2")
    para(doc, "ชื่อหัวคอลัมน์ในรายงานเรียงตามลำดับรหัส คือ 0 = ปกติ, 1 = เลือดออก, 2 = เหงือกอักเสบ, "
              "3 = ปริทันต์, 9 = ตรวจไม่ได้", size=14, color=GREY,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=3)

    # ---------- 4 ----------
    doc.add_page_break()
    h1(doc, "๔. เงื่อนไขการคำนวณรายตาราง")
    para(doc, "ทุกตารางตั้งแต่ตารางที่ 2 เป็นต้นไป ใช้ตัวหารเป็นจำนวนผู้ผ่านเกณฑ์คุณภาพของกลุ่มอายุนั้น "
              "ยกเว้นที่ระบุไว้เป็นอย่างอื่น", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    h2(doc, "ตารางที่ ๑ คุณภาพของข้อมูลการตรวจ")
    table(doc, [
        ["ช่อง", "วิธีคำนวณ"],
        ["จำนวนที่ตรวจ", "นับผู้ที่ `date_serv` ไม่ว่าง และอยู่ในปีงบประมาณที่วิเคราะห์"],
        ["จำนวนที่ผ่านเกณฑ์คุณภาพ", "นับผู้ที่ผ่านเงื่อนไขในหัวข้อ ๒ ชั้นที่ ๓"],
        ["ร้อยละ", "จำนวนที่ผ่านเกณฑ์ ÷ จำนวนที่ตรวจ × 100"],
    ], [4.5, 12.0], sizes=13, aligns=["l", "l"])

    h2(doc, "ตารางที่ ๒ ร้อยละของผู้ปราศจากโรคฟันผุและมีประสบการณ์ฟันผุ")
    para(doc, "ตัวหาร คือ จำนวนผู้ผ่านเกณฑ์คุณภาพ · ตัวตั้ง คือ จำนวนคน (ไม่ใช่จำนวนซี่) ที่เข้าเงื่อนไข "
              "โดย c, f, x หมายถึงฟันผุ ฟันอุด และฟันถอน ของชุดฟันที่ใช้ในกลุ่มอายุนั้น",
         size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    table(doc, [
        ["คอลัมน์", "เงื่อนไข (อายุ 3 ปี ใช้ฟันน้ำนม)", "เงื่อนไข (อายุ 6, 12, 60+ ใช้ฟันแท้)"],
        ["ปราศจากฟันผุ", "`dcaries + dfilling + dextract = 0`", "`pcaries + pfilling + pextract = 0`"],
        ["ฟันผุไม่ได้รักษา", "`dcaries > 0`", "`pcaries > 0`"],
        ["ฟันถอน", "`dextract > 0`", "`pextract > 0`"],
        ["ฟันอุด", "`dfilling > 0`", "`pfilling > 0`"],
        ["ฟันผุถอนอุด", "`dcaries + dfilling + dextract > 0`", "`pcaries + pfilling + pextract > 0`"],
    ], [3.6, 6.4, 6.5], sizes=13, aligns=["l", "l", "l"])

    h2(doc, "ตารางที่ ๓ ค่าเฉลี่ยฟันผุ ถอน อุด (ซี่ต่อคน)")
    table(doc, [
        ["คอลัมน์", "วิธีคำนวณ (N = จำนวนผู้ผ่านเกณฑ์คุณภาพ)"],
        ["ฟันที่มีในปาก", "ผลรวม `dteeth` หรือ `pteeth` ÷ N"],
        ["ฟันผุ", "ผลรวม `dcaries` หรือ `pcaries` ÷ N"],
        ["ฟันถอน", "ผลรวม `dextract` หรือ `pextract` ÷ N"],
        ["ฟันอุด", "ผลรวม `dfilling` หรือ `pfilling` ÷ N"],
        ["ฟันผุถอนอุด", "ผลรวมของสามช่องข้างต้น (ค่า dmft สำหรับอายุ 3 ปี และ DMFT สำหรับอายุอื่น)"],
    ], [4.0, 12.5], sizes=13, aligns=["l", "l"])

    h2(doc, "ตารางที่ ๔ ร้อยละของผู้มีภาวะเหงือกอักเสบและสภาวะปริทันต์")
    bullet(doc, "ใช้เฉพาะกลุ่มอายุ 12 ปี และ 60 ปีขึ้นไป")
    bullet(doc, "ตัวหาร คือ จำนวนผู้ผ่านเกณฑ์คุณภาพที่มีค่า gum ใช้งานได้ (แสดงในคอลัมน์ “รวม” ท้ายตาราง)")
    bullet(doc, "ตัวตั้ง คือ จำนวนคนในแต่ละค่าของ GumStatus ตามหัวข้อ ๓.๒")

    h2(doc, "ตารางที่ ๕ ค่าเฉลี่ยจำนวน sextant ในแต่ละสภาวะ")
    bullet(doc, "ใช้เฉพาะกลุ่มอายุ 12 ปี และ 60 ปีขึ้นไป และใช้ฐานเดียวกับตารางที่ 4")
    bullet(doc, "นับจำนวน sextant ของแต่ละคนแยกตามค่า 0, 1, 2, 3 และ 9 (รวมกันได้ 6 ต่อคน) "
                "แล้วหารด้วยจำนวนคน จึงได้ค่าเฉลี่ยจำนวน sextant ต่อคนในแต่ละสภาวะ")

    h2(doc, "ตารางที่ ๖ ผู้สูงอายุที่มีฟันใช้งานและคู่สบฟันหลัง")
    bullet(doc, "ใช้เฉพาะกลุ่มอายุ 60 ปีขึ้นไป ตัวหาร คือ จำนวนผู้ผ่านเกณฑ์คุณภาพ")
    table(doc, [
        ["คอลัมน์", "เงื่อนไข"],
        ["ฟันใช้งาน 20 ซี่", "`pteeth ≥ 20`"],
        ["คู่สบฟันหลัง 4 คู่ขึ้นไป", "`permanent_permanent + permanent_prosthesis + prosthesis_prosthesis ≥ 4`"],
        ["ทั้งสองเงื่อนไข", "เข้าเงื่อนไขทั้งสองข้อพร้อมกัน"],
    ], [4.5, 12.0], sizes=13, aligns=["l", "l"])

    h2(doc, "ตารางที่ ๗ ความจำเป็นในการรักษา")
    para(doc, "ตัวหาร คือ จำนวนผู้ผ่านเกณฑ์คุณภาพ · ตัวตั้ง คือ จำนวนคนที่เข้าเงื่อนไข",
         size=14, color=GREY)
    table(doc, [
        ["คอลัมน์", "เงื่อนไข"],
        ["ทาฟลูออไรด์", "`need_fluoride = 1`"],
        ["เคลือบหลุมร่องฟัน", "`need_sealant > 0`"],
        ["อุดฟัน", "อายุ 3 ปี ใช้ `need_dfilling > 0` · อายุ 6, 12 และ 60+ ใช้ `need_pfilling > 0`"],
        ["ถอนฟัน / รักษารากฟัน", "อายุ 3 ปี ใช้ `need_dextract > 0` · อายุ 6, 12 และ 60+ ใช้ `need_pextract > 0`"],
        ["ขูดหินปูน", "`need_scaling = 1`"],
    ], [4.5, 12.0], sizes=13, aligns=["l", "l"])

    # ---------- 5 ----------
    h1(doc, "๕. จุดที่ตีความเพิ่มเติมจากเอกสารต้นฉบับ")
    para(doc, "เอกสารเงื่อนไขต้นฉบับมีบางจุดที่กำกวมหรือพิมพ์คลาดเคลื่อน จึงตีความไว้ดังนี้ "
              "และได้รับการยืนยันจากผู้รับผิดชอบงานแล้ว", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    table(doc, [
        ["ประเด็น", "ข้อความในเอกสารต้นฉบับ", "สิ่งที่ระบบใช้จริง"],
        ["ตารางที่ 3 อายุ 3 ปี", "ระบุค่าเฉลี่ยฟันถอนและฟันอุดว่าใช้ dcaries ทั้งคู่",
         "ใช้ `dextract` สำหรับฟันถอน และ `dfilling` สำหรับฟันอุด"],
        ["ตารางที่ 6", "ระบุ “นับจำนวนคนที่มี PTEETH = 20”", "ใช้ `pteeth ≥ 20` ตามนิยาม functional dentition"],
        ["ตารางที่ 7", "คำว่า “เคลือบหลุมร่องฟัน” ปรากฏซ้ำสามบรรทัด",
         "เคลือบหลุมร่องฟันใช้ `need_sealant > 0` ส่วนอุดฟันและถอนฟันจับคู่ตามหัวตารางจริง"],
        ["ค่า gum ผิดรูปแบบ", "ไม่ได้ระบุวิธีจัดการ",
         "ค่าสั้นกว่า 6 หลักถือเป็นค่าเดียวทั้งปาก และเลข 4 กับ 5 นับเป็น 3"],
        ["การแบ่งกลุ่มอายุ", "ไม่ได้ระบุ", "ใช้ชื่อโฟลเดอร์ (คอลัมน์ agesurvey) ไม่ใช้ `denttype` หรือ `age_y`"],
    ], [3.4, 6.6, 6.5], sizes=13, aligns=["l", "l", "l"])

    # ---------- 6 ----------
    h1(doc, "๖. การนำไปใช้และการปรับปรุงข้อมูล")
    bullet(doc, "โปรแกรมแจกจ่ายเป็นไฟล์ .exe ไฟล์เดียว ผู้ใช้ดาวน์โหลดจาก GitHub นำไฟล์ Excel ของพื้นที่ตนเอง "
                "มาวางในโฟลเดอร์ data ตามกลุ่มอายุ แล้วดับเบิลคลิกโปรแกรม")
    bullet(doc, "เปิดโปรแกรมครั้งแรกจะสร้างโฟลเดอร์ data/3ปี, data/6ปี, data/12ปี และ data/60ปี ให้อัตโนมัติ")
    bullet(doc, "เมื่อประมวลผลเสร็จ ระบบจะเปิดไฟล์ ผลลัพธ์\\dashboard.html ในเบราว์เซอร์ให้เอง "
                "ไฟล์นี้เปิดซ้ำได้ตลอดโดยไม่ต้องเชื่อมต่ออินเทอร์เน็ต")
    bullet(doc, "ข้อมูลที่ฝังอยู่ในหน้าเว็บเป็นตัวเลขสรุประดับหน่วยบริการเท่านั้น ไม่มีข้อมูลรายบุคคล "
                "การกรองระดับจังหวัด อำเภอ และหน่วยบริการ ทำได้โดยรวมตัวเลขของหน่วยบริการที่เลือก")
    bullet(doc, "ไฟล์ combined.csv ที่โปรแกรมสร้างระหว่างทาง ตัดคอลัมน์ชื่อ สกุล เลขบัตรประชาชน ที่อยู่ "
                "วันเกิด และรหัสหมู่บ้าน ออกแล้ว หากจำเป็นต้องตรวจสอบรายบุคคล ใช้ตัวเลือก --full-csv")
    bullet(doc, "ครั้งต่อไปโปรแกรมจะข้ามการอ่าน Excel ซ้ำถ้าไฟล์ไม่เปลี่ยนแปลง หากต้องการบังคับให้อ่านใหม่ "
                "ใช้ตัวเลือก --force หรือลบไฟล์ ผลลัพธ์\\combined.csv")
    bullet(doc, "ปีงบประมาณจะถูกคำนวณใหม่อัตโนมัติทุกครั้งจากวันที่ตรวจล่าสุดในข้อมูลชุดนั้น")

    para(doc, "ไฟล์ในโครงการ", bold=True, space_before=8, space_after=3)
    table(doc, [
        ["ไฟล์", "หน้าที่"],
        ["`src/build_dashboard.py`", "โค้ดหลัก รวมไฟล์ Excel คำนวณทุกตาราง และสร้างรายงาน"],
        ["`src/dashboard_template.html`", "แม่แบบหน้าเว็บ (ส่วนแสดงผลและตัวกรอง)"],
        ["`src/hospitals.csv`", "ตารางเทียบรหัสหน่วยบริการกับอำเภอและจังหวัด (เข้ารหัส cp874)"],
        ["`ผลลัพธ์/dashboard.html`", "รายงานที่ใช้งานจริง ฝังข้อมูลสรุปไว้ภายใน"],
        ["`ผลลัพธ์/combined.csv`", "ข้อมูลทุกไฟล์ต่อกันพร้อมคอลัมน์ agesurvey สำหรับตรวจสอบย้อนหลัง"],
        ["`ผลลัพธ์/summary.json`", "ตัวเลขสรุประดับหน่วยบริการที่ใช้ฝังในหน้าเว็บ"],
        ["`ผลลัพธ์/log.txt`", "บันทึกการทำงานและข้อผิดพลาด สำหรับส่งให้ผู้ดูแลระบบเมื่อมีปัญหา"],
    ], [5.6, 10.9], sizes=13, aligns=["l", "l"])

    para(doc, f"เอกสารฉบับนี้อ้างอิงผลการประมวลผลข้อมูลปีงบประมาณ {s['fiscalYear']} "
              f"จำนวน {f(s['rowsTotal'])} รายการ จาก {len(s['units'])} หน่วยบริการ "
              f"จัดทำเมื่อ {s['generatedAt']}",
         size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=14)

    out = DOCS / "เอกสารอธิบายเงื่อนไขการคำนวณ.docx"
    doc.save(out)
    print("saved:", out)
    return out


if __name__ == "__main__":
    build()
